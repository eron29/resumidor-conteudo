import re

import pandas as pd
import streamlit as st
from docx import Document
from pypdf import PdfReader
from youtube_transcript_api import (
    NoTranscriptFound,
    RequestBlocked,
    TranscriptsDisabled,
    YouTubeTranscriptApi,
)
from youtube_transcript_api.proxies import WebshareProxyConfig

YOUTUBE_ID_PATTERNS = [
    r"(?:v=|\/embed\/|\/shorts\/)([0-9A-Za-z_-]{11})",
    r"youtu\.be\/([0-9A-Za-z_-]{11})",
]


def extract_youtube_id(url: str) -> str:
    for pattern in YOUTUBE_ID_PATTERNS:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    raise ValueError("URL do YouTube inválida ou não reconhecida.")


def _get_secret(key: str):
    try:
        return st.secrets.get(key)
    except Exception:
        return None


def _build_youtube_api() -> YouTubeTranscriptApi:
    proxy_username = _get_secret("WEBSHARE_PROXY_USERNAME")
    proxy_password = _get_secret("WEBSHARE_PROXY_PASSWORD")
    if proxy_username and proxy_password:
        return YouTubeTranscriptApi(
            proxy_config=WebshareProxyConfig(
                proxy_username=proxy_username,
                proxy_password=proxy_password,
            )
        )
    return YouTubeTranscriptApi()


def extract_from_youtube(url: str, languages=("pt", "pt-BR", "en")) -> str:
    video_id = extract_youtube_id(url)
    api = _build_youtube_api()
    try:
        transcript = api.fetch(video_id, languages=list(languages))
    except NoTranscriptFound:
        # Nenhuma legenda nos idiomas preferidos: usa a primeira disponível, seja qual for o idioma.
        try:
            transcript = next(iter(api.list(video_id))).fetch()
        except (TranscriptsDisabled, NoTranscriptFound, StopIteration) as exc:
            raise RuntimeError("Este vídeo não possui legendas/transcrição disponível.") from exc
    except TranscriptsDisabled as exc:
        raise RuntimeError("Este vídeo não possui legendas/transcrição disponível.") from exc
    except RequestBlocked as exc:
        raise RuntimeError(
            "O YouTube está bloqueando as requisições vindas do servidor deste app (comum em "
            "hospedagens na nuvem, como o Streamlit Community Cloud). Para resolver, configure um "
            "proxy gratuito: crie uma conta em webshare.io, copie as credenciais de 'Proxy' e "
            "adicione WEBSHARE_PROXY_USERNAME e WEBSHARE_PROXY_PASSWORD nos Secrets do app."
        ) from exc
    return " ".join(snippet.text for snippet in transcript)


def extract_from_pdf(file) -> str:
    reader = PdfReader(file)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def extract_from_docx(file) -> str:
    doc = Document(file)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_from_text(file) -> str:
    raw = file.read()
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="ignore")
    return raw


def extract_from_spreadsheet(file, filename: str) -> str:
    if filename.lower().endswith(".csv"):
        df = pd.read_csv(file)
        return df.to_markdown(index=False)
    xls = pd.ExcelFile(file)
    parts = []
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        parts.append(f"### Planilha: {sheet}\n\n{df.to_markdown(index=False)}")
    return "\n\n".join(parts)


def extract_content(source_type: str, payload) -> str:
    """source_type: 'youtube' (payload = URL) ou 'arquivo' (payload = (filename, file-like))."""
    if source_type == "youtube":
        return extract_from_youtube(payload)

    filename, file = payload
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        return extract_from_pdf(file)
    if ext == "docx":
        return extract_from_docx(file)
    if ext in ("txt", "md"):
        return extract_from_text(file)
    if ext in ("xlsx", "csv"):
        return extract_from_spreadsheet(file, filename)

    raise ValueError(f"Formato de arquivo não suportado: .{ext}")
